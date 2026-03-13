from docx import Document as DocxDocument
from docx.shared import Pt, Inches
import datetime


def generate_word_report(entity, analysis, docs, output_path: str):
    doc = DocxDocument()

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_heading("IntelliCredit AI", 0)
    doc.add_heading("Credit Assessment Report", 1)
    doc.add_paragraph(f"Entity: {entity.company_name}")
    doc.add_paragraph(f"CIN: {entity.cin or 'N/A'}  |  PAN: {entity.pan or 'N/A'}")
    doc.add_paragraph(f"Sector: {entity.sector or 'N/A'}  |  Sub-sector: {entity.sub_sector or 'N/A'}")
    doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("CONFIDENTIAL — FOR INTERNAL USE ONLY")
    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(analysis.executive_summary or "")

    # ── Recommendation ─────────────────────────────────────────────────────
    doc.add_heading("Loan Recommendation", 1)
    rec_table = doc.add_table(rows=2, cols=5)
    rec_table.style = "Table Grid"
    headers = ["Decision", "Grade", "Recommended Limit", "Interest Rate", "Tenure"]
    values = [
        analysis.recommendation or "",
        analysis.grade or "",
        f"Rs {analysis.recommended_limit_cr} Cr" if analysis.recommended_limit_cr else "",
        f"{analysis.recommended_rate_pct}%" if analysis.recommended_rate_pct else "",
        f"{analysis.recommended_tenure_months} months" if analysis.recommended_tenure_months else "",
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        rec_table.cell(0, i).text = h
        rec_table.cell(1, i).text = v

    if analysis.conditions:
        doc.add_heading("Conditions", 2)
        for c in analysis.conditions:
            doc.add_paragraph(f"• {c}")

    # ── Score Breakdown ────────────────────────────────────────────────────
    doc.add_heading("Score Breakdown", 1)
    if analysis.scores:
        score_table = doc.add_table(rows=1, cols=2)
        score_table.style = "Table Grid"
        score_table.cell(0, 0).text = "Dimension"
        score_table.cell(0, 1).text = "Score /100"
        for k, v in analysis.scores.items():
            row = score_table.add_row()
            row.cells[0].text = k.replace("_", " ").title()
            row.cells[1].text = str(v)

    # ── Reasoning Chain ────────────────────────────────────────────────────
    doc.add_heading("Reasoning Engine", 1)
    if analysis.reasoning_chain:
        for step in analysis.reasoning_chain:
            doc.add_heading(f"Step {step.get('step')}: {step.get('factor')}", 3)
            doc.add_paragraph(
                f"Signal: {step.get('signal')}  |  Impact: {step.get('impact')}  |  Weight: {step.get('weight')}"
            )
            doc.add_paragraph(step.get("evidence", ""))

    # ── SWOT ───────────────────────────────────────────────────────────────
    doc.add_heading("SWOT Analysis", 1)
    if analysis.swot:
        for quadrant, items in analysis.swot.items():
            doc.add_heading(quadrant.title(), 2)
            for item in (items or []):
                doc.add_paragraph(f"• {item}")

    # ── Secondary Research ─────────────────────────────────────────────────
    doc.add_heading("Secondary Research — Tavily", 1)
    if analysis.sentiment:
        s = analysis.sentiment
        doc.add_paragraph(
            f"Overall Sentiment: {s.get('overall_sentiment')} (score: {s.get('sentiment_score')})"
        )
        if s.get("red_flags"):
            doc.add_heading("Red Flags", 2)
            for flag in s["red_flags"]:
                doc.add_paragraph(f"[!] {flag}")
        if s.get("sector_outlook"):
            doc.add_heading("Sector Outlook", 2)
            doc.add_paragraph(s["sector_outlook"])

    doc.add_heading("News Articles", 2)
    for article in (analysis.news_results or [])[:5]:
        doc.add_paragraph(f"• {article.get('title')}")
        doc.add_paragraph(f"  {article.get('url')}", style="Normal")

    # ── Triangulation ──────────────────────────────────────────────────────
    doc.add_heading("Cross-Document Triangulation", 1)
    if analysis.triangulation:
        for k, v in analysis.triangulation.items():
            if k != "inconsistencies":
                doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        if analysis.triangulation.get("inconsistencies"):
            doc.add_heading("Inconsistencies Found", 2)
            for inc in analysis.triangulation["inconsistencies"]:
                doc.add_paragraph(f"[!] {inc}")

    # ── Risks & Mitigants ──────────────────────────────────────────────────
    doc.add_heading("Key Risks & Mitigants", 1)
    if analysis.key_risks:
        doc.add_heading("Key Risks", 2)
        for r in analysis.key_risks:
            doc.add_paragraph(f"• {r}")
    if analysis.mitigants:
        doc.add_heading("Mitigants", 2)
        for m in analysis.mitigants:
            doc.add_paragraph(f"• {m}")

    # ── Appendix: Extracted Data ───────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix — Extracted Document Data", 1)
    for d in docs:
        if d.extracted_data and d.schema_status == "done":
            doc.add_heading(f"{d.confirmed_doc_type}: {d.filename}", 2)
            for k, v in d.extracted_data.items():
                if not k.startswith("_") and v is not None:
                    doc.add_paragraph(f"{k}: {v}")

    doc.save(output_path)
